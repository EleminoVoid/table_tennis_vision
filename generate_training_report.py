import argparse
import datetime as dt
import os
import re
from pathlib import Path
from typing import Optional


def read_text(path: Path) -> str:
    if not path.exists():
        return ""
    return path.read_text(encoding="utf-8", errors="ignore")


def extract_train_block(script_text: str) -> str:
    marker = "results = model.train("
    start = script_text.find(marker)
    if start == -1:
        return "Training block not found."

    depth = 0
    end = None
    for index in range(start, len(script_text)):
        char = script_text[index]
        if char == "(":
            depth += 1
        elif char == ")":
            depth -= 1
            if depth == 0:
                end = index + 1
                break

    if end is None:
        return script_text[start:start + 4000]
    return script_text[start:end]


def find_latest_log(repo_root: Path) -> Optional[Path]:
    candidates = []
    patterns = [
        "training_log_safe_*.txt",
        "training_log_cuda_*.txt",
        "training_log*.txt",
        "training_err*.txt",
    ]
    for pattern in patterns:
        candidates.extend(repo_root.glob(pattern))

    candidates = [path for path in candidates if path.is_file()]
    if not candidates:
        return None
    return max(candidates, key=lambda path: path.stat().st_mtime)


def build_prompt(
    workflow_text: str,
    train_block: str,
    log_file_name: str,
    log_tail: str,
    log_chars_used: int,
) -> str:
    return f"""
You are preparing a professional training report for a computer vision project.
Use the project context and log data below.

Requirements:
- Write clear, concise, technical English.
- Include these sections with headings exactly:
  1) Executive Summary
  2) Training Configuration
  3) Runtime Stability Notes
  4) Observed Progress & Signals
  5) Risks / Unknowns
  6) Recommended Next Steps
- In Training Configuration, include key parameters (device, batch, workers, cache, epochs, save behavior).
- In Runtime Stability Notes, call out signs of interruption risk and mitigation already applied.
- In Recommended Next Steps, provide practical, actionable bullets.
- Do not invent metrics that are not in the provided data.

Project workflow context:
{workflow_text}

Current model.train(...) block:
{train_block}

Training log source: {log_file_name}
Log tail used (last {log_chars_used} characters):
{log_tail}
""".strip()


def _find_setting(train_block: str, key: str) -> str:
    match = re.search(rf"\b{re.escape(key)}\s*=\s*([^,\n]+)", train_block)
    return match.group(1).strip() if match else "N/A"


def build_local_report(
    train_block: str,
    log_tail: str,
    log_file_name: str,
    fallback_reason: str = "",
) -> str:
    lines = [line.strip() for line in log_tail.splitlines() if line.strip()]
    warning_lines = [line for line in lines if "warning" in line.lower()]
    error_lines = [line for line in lines if "error" in line.lower() or "traceback" in line.lower()]
    epoch_lines = [line for line in lines if re.search(r"\b\d+/\d+\b", line)]
    epoch_preview = epoch_lines[-5:] if epoch_lines else []

    device = _find_setting(train_block, "device")
    batch = _find_setting(train_block, "batch")
    workers = _find_setting(train_block, "workers")
    cache = _find_setting(train_block, "cache")
    epochs = _find_setting(train_block, "epochs")
    save_period = _find_setting(train_block, "save_period")

    risks = []
    if warning_lines:
        risks.append("Warnings are present in the training log and should be reviewed.")
    if error_lines:
        risks.append("Error-like lines were found in the log tail and may indicate instability.")
    if not error_lines and not warning_lines:
        risks.append("No obvious errors detected in the provided log tail.")

    if fallback_reason:
        summary_note = f"Local report mode was used: {fallback_reason}."
    else:
        summary_note = "Local report mode was used (no paid API required)."

    body = [
        "# Executive Summary",
        summary_note,
        f"This report summarizes training behavior using local log parsing from `{log_file_name}`.",
        "",
        "# Training Configuration",
        f"- device: {device}",
        f"- batch: {batch}",
        f"- workers: {workers}",
        f"- cache: {cache}",
        f"- epochs: {epochs}",
        f"- save_period: {save_period}",
        "",
        "# Runtime Stability Notes",
        "- This local report does not call external APIs, so it cannot fail from credit limits.",
        "- Stability still depends on GPU/CPU resources and training configuration.",
        "",
        "# Observed Progress & Signals",
    ]

    if epoch_preview:
        body.append("- Recent epoch/progress lines:")
        body.extend([f"- {line}" for line in epoch_preview])
    else:
        body.append("- No epoch progress lines were detected in the provided log tail.")

    if warning_lines:
        body.append("- Recent warnings:")
        body.extend([f"- {line}" for line in warning_lines[-5:]])

    if error_lines:
        body.append("- Recent error-like lines:")
        body.extend([f"- {line}" for line in error_lines[-5:]])

    body.extend([
        "",
        "# Risks / Unknowns",
    ])
    body.extend([f"- {item}" for item in risks])

    body.extend([
        "",
        "# Recommended Next Steps",
        "- Continue training and monitor for new warnings/errors.",
        "- If interruptions continue, reduce batch size and dataloader workers.",
        "- Keep per-epoch checkpoints enabled for recovery.",
        "- Re-run this report after each major training session.",
    ])

    return "\n".join(body).strip()


def call_claude(api_key: str, model: str, prompt: str) -> str:
    try:
        from anthropic import Anthropic
    except ImportError as exc:
        raise RuntimeError(
            "Missing dependency 'anthropic'. Install with: pip install anthropic"
        ) from exc

    client = Anthropic(api_key=api_key)
    response = client.messages.create(
        model=model,
        max_tokens=2800,
        temperature=0.2,
        messages=[{"role": "user", "content": prompt}],
    )

    chunks = []
    for block in response.content:
        text = getattr(block, "text", None)
        if text:
            chunks.append(text)
    return "\n".join(chunks).strip()


def write_docx(output_path: Path, title: str, body_text: str, log_name: str, model: str) -> None:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "Missing dependency 'python-docx'. Install with: pip install python-docx"
        ) from exc

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"Generated: {dt.datetime.now().isoformat(timespec='seconds')}")
    doc.add_paragraph(f"Claude model: {model}")
    doc.add_paragraph(f"Log source: {log_name}")

    for raw_line in body_text.splitlines():
        line = raw_line.strip()
        if not line:
            doc.add_paragraph("")
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line[:3].isdigit() and line[1:3] == ") ":
            doc.add_paragraph(line, style="List Number")
        else:
            doc.add_paragraph(line)

    doc.save(output_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Generate a Word (.docx) training report using Claude API or local mode"
    )
    parser.add_argument(
        "--log-file",
        type=str,
        default="",
        help="Path to training log file (default: auto-detect latest training log)",
    )
    parser.add_argument(
        "--output",
        type=str,
        default="",
        help="Output .docx path (default: reports/training_report_<timestamp>.docx)",
    )
    parser.add_argument(
        "--model",
        type=str,
        default="claude-3-7-sonnet-latest",
        help="Anthropic model ID",
    )
    parser.add_argument(
        "--max-log-chars",
        type=int,
        default=120000,
        help="How many trailing log characters to include in the prompt",
    )
    parser.add_argument(
        "--api-key",
        type=str,
        default="",
        help="Anthropic API key (default: ANTHROPIC_API_KEY env var)",
    )
    parser.add_argument(
        "--local-only",
        action="store_true",
        help="Skip Claude API and build the report locally from logs/config",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    repo_root = Path(__file__).resolve().parent

    api_key = args.api_key or os.getenv("ANTHROPIC_API_KEY", "")

    script_path = repo_root / "download_and_train.py"
    workflow_path = repo_root / "TRAINING_WORKFLOW.md"

    if args.log_file:
        log_path = Path(args.log_file)
        if not log_path.is_absolute():
            log_path = repo_root / log_path
    else:
        found = find_latest_log(repo_root)
        if found is None:
            raise SystemExit("No training log file found. Pass --log-file explicitly.")
        log_path = found

    if not log_path.exists():
        raise SystemExit(f"Log file does not exist: {log_path}")

    script_text = read_text(script_path)
    workflow_text = read_text(workflow_path)
    log_text = read_text(log_path)

    if not script_text:
        raise SystemExit("Could not read download_and_train.py")

    train_block = extract_train_block(script_text)
    max_chars = max(5000, args.max_log_chars)
    log_tail = log_text[-max_chars:]

    prompt = build_prompt(
        workflow_text=workflow_text,
        train_block=train_block,
        log_file_name=log_path.name,
        log_tail=log_tail,
        log_chars_used=len(log_tail),
    )

    report_model_label = args.model
    report_text = ""

    if args.local_only:
        report_model_label = "local-offline"
        report_text = build_local_report(
            train_block=train_block,
            log_tail=log_tail,
            log_file_name=log_path.name,
            fallback_reason="--local-only flag",
        )
    elif not api_key:
        report_model_label = "local-offline"
        report_text = build_local_report(
            train_block=train_block,
            log_tail=log_tail,
            log_file_name=log_path.name,
            fallback_reason="missing API key",
        )
    else:
        try:
            report_text = call_claude(api_key=api_key, model=args.model, prompt=prompt)
            if not report_text:
                raise RuntimeError("Claude returned an empty response")
        except Exception as exc:
            report_model_label = "local-offline"
            report_text = build_local_report(
                train_block=train_block,
                log_tail=log_tail,
                log_file_name=log_path.name,
                fallback_reason=f"Claude unavailable ({exc})",
            )

    if args.output:
        output_path = Path(args.output)
        if not output_path.is_absolute():
            output_path = repo_root / output_path
    else:
        timestamp = dt.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = repo_root / "reports" / f"training_report_{timestamp}.docx"

    write_docx(
        output_path=output_path,
        title="Table Tennis Vision Training Report",
        body_text=report_text,
        log_name=log_path.name,
        model=report_model_label,
    )

    print(f"Report written to: {output_path}")


if __name__ == "__main__":
    main()
